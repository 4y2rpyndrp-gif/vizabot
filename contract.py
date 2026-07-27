"""
Mijoz ma'lumotlari asosida TO'LIQ (qisqartirilmagan) shartnoma DOCX faylini
avtomatik yaratadi. python-docx kutubxonasidan foydalanadi (Railway
konteynerida LibreOffice kerak emas). Matn asl "Shartnoma_yangilangan_
kuchaytirilgan.docx" andozasiga asoslangan, faqat mijozga xos joylar
(ism, sana, pasport, narx) avtomatik to'ldiriladi.
"""

import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from pricing import get_country_pricing

COMPANY_NAME = "«AUROVIA» MCHJ"
COMPANY_DIRECTOR = "Buriyev Sh."
COMPANY_STIR = "312 742 107"
COMPANY_MFO = "01158"
COMPANY_ACCOUNT = "2020 8000 9073 9633 0001"
COMPANY_PHONE = "+998 78 122 00 66"
COMPANY_ADDRESS = "Samarqand shahri, M. Ulug'bek ko'chasi 34"


# ---------- Yordamchi funksiyalar ----------

def _heading(doc, text, size=13, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    return p


def _body(doc, text, bold=False, size=11, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p


def _bullet(doc, text, size=11):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("•  " + text)
    run.font.size = Pt(size)
    return p


def _notice_box(doc, text):
    """3-bandning muhim shartini ajratib ko'rsatish uchun ramkali jadval."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.width = Cm(16)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


# ---------- Asosiy funksiya ----------

def generate_contract_docx(
    contract_number: str,
    full_name: str,
    birth_date: str,
    passport: str,
    phone: str,
    address: str,
    country: str,
    output_dir: str = "/tmp",
) -> str:
    pricing = get_country_pricing(country) or {"total": 0, "prepay": 0, "remainder": 0}
    today = datetime.now().strftime("%d.%m.%Y")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    # ---------- Sarlavha ----------
    _heading(doc, "VIZA OLISHGA KO'MAKLASHISH KONSALTING XIZMATLARINI KO'RSATISH BO'YICHA",
             size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    _heading(doc, f"SHARTNOMA № {contract_number}", size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    _body(doc,
        f"{COMPANY_NAME} (keyingi o'rinlarda — «Ijrochi») nomidan Direktor {COMPANY_DIRECTOR} "
        f"shaxsida, bir tomondan, va {full_name} (keyingi o'rinlarda — «Buyurtmachi») shaxsida "
        f"ikkinchi tomondan, birgalikda «Tomonlar», alohida-alohida «Tomon» deb ataluvchi ushbu "
        f"Shartnomani quyidagilar to'g'risida tuzdilar:")

    # ---------- 1. Atamalar ----------
    _heading(doc, "1. ATAMALAR VA UMUMIY QOIDALAR")
    _bullet(doc, "«Xizmatlar» deganda Ijrochining Buyurtmachiga xorijiy davlatda ishga "
                 "joylashish va tegishli ish vizasini rasmiylashtirish jarayonida "
                 "ko'rsatadigan konsalting, tashkiliy va maslahat xizmatlari tushuniladi.")
    _bullet(doc, "«Dalolatnoma» deganda Tomonlar tomonidan imzolanadigan, muayyan xizmat "
                 "bosqichi bajarilganligini tasdiqlovchi ikki tomonlama hujjat tushuniladi "
                 "(ushbu Shartnomaning 1-ILOVASI).")
    _bullet(doc, "«Ish beruvchi» va «Elchixona» — Buyurtmachi ishga joylashmoqchi bo'lgan "
                 "xorijiy tashkilot va tegishli davlatning O'zbekistondagi yoki mintaqadagi "
                 "diplomatik vakolatxonasi, ular Ijrochiga qaram bo'lmagan mustaqil uchinchi "
                 "shaxslar hisoblanadi.")
    _body(doc, "Shartnoma shartlari O'zbekiston Respublikasining amaldagi qonunchiligiga, "
                "jumladan Fuqarolik kodeksining pudrat va xizmat ko'rsatish shartnomalariga "
                "oid qoidalariga asoslanadi. Shartnoma imzolangan kundan e'tiboran har ikkala "
                "Tomon uning barcha shartlari bilan to'liq tanishgan va ularga so'zsiz rozi "
                "bo'lgan hisoblanadi.")

    # ---------- 2. Predmet ----------
    _heading(doc, "2. SHARTNOMA PREDMETI VA XIZMATLAR TAVSIFI")
    _body(doc, "2.1. Ijrochi Buyurtmachiga quyidagi xizmatlarni ko'rsatishni o'z zimmasiga oladi:")
    _bullet(doc, "Ish vizasini olish uchun zarur bo'lgan hujjatlar ro'yxatini tuzish va "
                 "ularni to'g'ri rasmiylashtirish bo'yicha yo'riqnoma berish;")
    _bullet(doc, "Buyurtmachining kasbiy tajribasiga mos professional rezyume shakllantirish "
                 "va uni potentsial ish beruvchilarga taqdim etish;")
    _bullet(doc, "Ish beruvchi Buyurtmachi nomzodini ko'rib chiqqan taqdirda, u bilan onlayn "
                 "suhbat (intervyu) tashkil etish va Buyurtmachini shu suhbatga tayyorlash;")
    _bullet(doc, "Ish beruvchi bilan muzokaralar muvaffaqiyatli yakunlanib, mehnat shartnomasi "
                 "imzolangandan so'ng — elchixonaga topshiriladigan hujjatlar to'plamini "
                 "shakllantirish va anketa blankalarini to'ldirish;")
    _bullet(doc, "Elchixonadagi viza suhbatiga nazariy va amaliy tayyorgarlik ko'rish.")
    _body(doc, "2.2. Taxminiy ish jadvali: ish beruvchi tomonidan hujjatlarni ko'rib chiqish — "
                "3 (uch) oygacha; elchixona tomonidan ariza ko'rib chiqilishi va suhbat — 1 "
                "(bir) oygacha; jarayonning umumiy belgilangan muddati — 4 (to'rt) oy. Ushbu "
                "muddatlar taxminiy bo'lib, ular to'liq yoki qisman ish beruvchi va "
                "elchixonaning ichki ish yuritish tartibiga bog'liq.")

    # ---------- 3-band (ramkali) ----------
    _heading(doc, "3-BAND. NATIJANING KAFOLATLANMASLIGI TO'G'RISIDA MUHIM SHART", size=12)
    _notice_box(doc,
        "Ijrochi Buyurtmachiga faqat KONSALTING VA TASHKILIY-TEXNIK XIZMATLAR ko'rsatadi — "
        "ya'ni jarayonni tayyorlaydi, hujjatlarni shakllantiradi, tarjima qiladi, suhbatga "
        "tayyorlaydi va tashkillashtiradi. Ijrochi ish beruvchining ishga qabul qilish qarori "
        "yoki elchixonaning viza berish/rad etish qaroriga na huquqiy, na amaliy jihatdan "
        "ta'sir ko'rsata olmaydi — bu qarorlar mutlaqo mustaqil uchinchi shaxslarning ichki "
        "vakolati doirasida qabul qilinadi.\n\n"
        "Shu sababli: (a) Buyurtmachining ish beruvchi bilan suhbatdan muvaffaqiyatli o'ta "
        "olmasligi, (b) elchixona tomonidan vizaning rad etilishi — Ijrochi tomonidan xizmat "
        "ko'rsatilmaganligi yoki sifatsiz ko'rsatilganligi degani EMAS va Ijrochining shartnoma "
        "bo'yicha javobgarligini keltirib chiqarmaydi, agar 5-bandda nazarda tutilgan "
        "Dalolatnomalar bilan tasdiqlangan tegishli xizmatlar belgilangan tartibda bajarilgan "
        "bo'lsa.\n\n"
        "Buyurtmachi ushbu Shartnomani imzolash bilan yuqoridagi shartni to'liq anglaganini "
        "va unga rozi ekanligini tasdiqlaydi.")

    # ---------- 4. To'lov ----------
    _heading(doc, "4. TO'LOV SHARTLARI")
    _body(doc, f"4.1. Ijrochi xizmatlarining umumiy qiymati {pricing['total']:,} so'mni tashkil qiladi.")
    _body(doc, "4.2. To'lov ikki bosqichda amalga oshiriladi:")
    _bullet(doc, f"Birinchi bosqich — shartnoma imzolangan kundan boshlab 3 (uch) bank ish "
                 f"kuni ichida {pricing['prepay']:,} so'm miqdorida boshlang'ich to'lov "
                 f"o'tkaziladi. Mazkur muddatda to'lov amalga oshirilmasa, Ijrochi Shartnomani "
                 f"bir tomonlama bekor qilish huquqini o'zida saqlab qoladi.")
    _bullet(doc, f"Ikkinchi bosqich — qolgan {pricing['remainder']:,} so'm summasi viza "
                 f"qo'lga kiritilgandan so'ng, kelishilgan jadval asosida to'lab boriladi.")
    _body(doc, "4.3. Boshlang'ich to'lov ushbu Shartnomaning 2.1-bandida sanab o'tilgan "
                "xizmatlarni amalga oshirish uchun zarur bo'lgan xarajatlarni (hujjatlarni "
                "tekshirish, tarjima qilish, rezyume tayyorlash, tashkiliy-texnik ishlar) "
                "qoplash maqsadida undiriladi. Ushbu xarajatlar xizmat ko'rsatilishi boshlangan "
                "zahoti yuzaga kelganligi sababli, boshlang'ich to'lov, ushbu Shartnomaning "
                "8-bandida ko'rsatilgan hollar bundan mustasno, qaytarilmas hisoblanadi.")

    # ---------- 5. Bajarish tartibi ----------
    _heading(doc, "5. XIZMATLARNI BAJARISH TARTIBI VA TASDIQLASH (DALOLATNOMALAR)")
    _body(doc, "5.1. Ijrochi tomonidan har bir xizmat bosqichi bajarilgach, Tomonlar ushbu "
                "Shartnomaning 1-ILOVASIDA keltirilgan shakldagi Dalolatnomani imzolaydilar. "
                "Dalolatnomada bajarilgan xizmat nomi, bajarilgan sanasi va Buyurtmachining "
                "ushbu xizmatni qabul qilib olganligi qayd etiladi.")
    _body(doc, "5.2. Dalolatnoma — ushbu Shartnoma bo'yicha xizmatning tegishli tartibda "
                "bajarilganligini tasdiqlovchi ASOSIY va YETARLI dalil hisoblanadi. Buyurtmachi "
                "Dalolatnomani imzolashdan asossiz bosh tortgan yoki imzolashdan bo'yin "
                "tovlagan taqdirda, xizmat Ijrochi tomonidan tegishli tarzda ko'rsatilgan, "
                "ammo Buyurtmachi tomonidan rasmiylashtirilmagan deb hisoblanadi; bunday "
                "holatda Ijrochi tomonidan yuborilgan yozma yoki elektron xabarnoma (SMS, "
                "messenjer, elektron pochta, qo'ng'iroqlar tarixi) xizmat ko'rsatilganligining "
                "muqobil dalili sifatida qabul qilinadi.")
    _body(doc, "5.3. Ijrochi xodimlarining Buyurtmachi bilan telefon orqali, messenjerlar "
                "orqali yoki boshqa aloqa vositalari orqali olib borgan barcha yozishmalari, "
                "qo'ng'iroqlar tarixi va suhbat vaqtini belgilash bo'yicha yuborilgan "
                "takliflar — xizmatning bajarilganligi yoki Buyurtmachining o'z majburiyatlarini "
                "bajarmaganligini (masalan, qo'ng'iroqqa javob bermaganligi, taklif etilgan "
                "suhbat vaqtiga kelmaganligi) tasdiqlovchi dalil sifatida ushbu Shartnomaga "
                "ilova qilinishi va har qanday davlat organi yoki sudga taqdim etilishi mumkin.")

    # ---------- 6. Aloqa ----------
    _heading(doc, "6. ALOQA, XABARNOMALAR VA YOZUVLARNI QAYD ETISH")
    _body(doc, "6.1. Tomonlar ushbu Shartnoma yuzasidan olib boriladigan telefon suhbatlari, "
                "onlayn uchrashuvlar va elektron yozishmalarning Ijrochi tomonidan xizmat "
                "ko'rsatish sifatini nazorat qilish va dalillashtirish maqsadida audio "
                "va/yoki video yozib olinishi, shuningdek skrinshot va boshqa elektron shaklda "
                "qayd etilishi mumkinligiga roziligini bildiradi.")
    _body(doc, "6.2. Ijrochi tomonidan Buyurtmachiga yuborilgan har qanday xabarnoma (suhbat "
                "kuni va vaqti, hujjat taqdim etish talabi va h.k.) yuborilgan sanadan "
                "e'tiboran tegishli tartibda yetkazilgan deb hisoblanadi, agar u Buyurtmachi "
                "tomonidan ko'rsatilgan aloqa vositalariga (telefon raqami, elektron pochta, "
                "messenjer) yuborilgan bo'lsa — bu Buyurtmachining xabarnomani real vaqtda "
                "ko'rgan yoki ko'rmaganligidan qat'i nazar amal qiladi.")
    _body(doc, "6.3. Buyurtmachi Shartnoma amal qilish davomida ko'rsatilgan telefon raqami "
                "va boshqa aloqa vositalari orqali doimiy ravishda bog'lanish imkoniyatini "
                "ta'minlashi, Ijrochi xodimlarining qo'ng'iroqlari va xabarlarini muntazam "
                "kuzatib borishi hamda ularga 24 (yigirma to'rt) soat ichida javob berishi "
                "shart. Aloqa vositasi vaqtincha ishlamay qolgan yoki o'zgargan taqdirda, "
                "Buyurtmachi bu haqda Ijrochini zudlik bilan xabardor qilishi lozim.")
    _body(doc, "6.4. Ijrochi Buyurtmachi bilan bog'lanish uchun kamida 2 (ikki) marta, turli "
                "aloqa vositalari (telefon qo'ng'irog'i va SMS yoki messenjer xabari) orqali "
                "harakat qiladi va bu urinishlarni sana, vaqt va natijasi ko'rsatilgan holda "
                "ushbu Shartnomaning 2-ILOVASIDA nazarda tutilgan Aloqa jurnaliga qayd etadi.")
    _body(doc, "6.5. Agar Buyurtmachi 6.4-bandda ko'rsatilgan tartibda amalga oshirilgan "
                "bog'lanish urinishlariga 6.3-bandda belgilangan muddatda javob bermasa, "
                "va/yoki unga aniq sana va vaqti ko'rsatilib taklif qilingan suhbat, uchrashuv "
                "yoki boshqa tadbirga o'z vaqtida kelmasa — tegishli xizmat (xabardor qilish, "
                "taklif etish, tashkillashtirish) Ijrochi tomonidan to'liq va tegishli tarzda "
                "bajarilgan deb hisoblanadi, Aloqa jurnalidagi yozuv esa buning yetarli dalili "
                "bo'ladi. Bunday holatda keyingi bosqichda yuzaga kelgan har qanday oqibat "
                "(jumladan, nomzodning ko'rib chiqilmasdan qolishi yoki suhbatdan "
                "chetlashtirilishi) faqat Buyurtmachining o'zi tomonidan bog'lanishga "
                "chiqmaganligi yoki uchrashuvga kelmaganligi natijasi hisoblanadi va Ijrochi "
                "javobgarligini keltirib chiqarmaydi.", bold=True)
    _body(doc, "6.6. Buyurtmachi ushbu Shartnomani imzolash bilan, agar u keyinchalik "
                "\"menga xabar berilmagan\" yoki \"men bilan bog'lanilmagan\" degan da'vo "
                "bilan chiqsa, ushbu da'voning to'g'riligi yoki noto'g'riligi birinchi "
                "navbatda 2-ILOVADAGI Aloqa jurnali va 5-banddagi Dalolatnomalar asosida "
                "aniqlanishiga roziligini bildiradi.")

    # ---------- 7. Majburiyatlar ----------
    _heading(doc, "7. TOMONLARNING MAJBURIYATLARI")
    _body(doc, "7.1. Ijrochi quyidagi majburiyatlarni o'z zimmasiga oladi:")
    _bullet(doc, "2.1-bandda sanab o'tilgan barcha xizmatlarni belgilangan muddatda va sifatli "
                 "ko'rsatish hamda buni Dalolatnomalar orqali rasmiylashtirish;")
    _bullet(doc, "Jarayonning har bir bosqichi to'g'risida Buyurtmachini o'z vaqtida xabardor "
                 "qilib borish;")
    _bullet(doc, "Buyurtmachi hujjatlari va shaxsiy ma'lumotlarining maxfiyligini ta'minlash.")
    _body(doc, "7.2. Buyurtmachi quyidagi majburiyatlarni bajarishi shart:")
    _bullet(doc, "So'ralgan hujjatlarni (biometrik pasport, ID karta, xorijga chiqish pasporti, "
                 "ta'lim diplomi, zarur bo'lsa haydovchilik guvohnomasi) belgilangan muddatda "
                 "taqdim etish;")
    _bullet(doc, "Shartnomada nazarda tutilgan to'lovlarni o'z vaqtida amalga oshirish;")
    _bullet(doc, "Onlayn intervyuda faol ishtirok etish va unga puxta tayyorgarlik ko'rish;")
    _bullet(doc, "Ijrochi xodimlarining qo'ng'iroqlari va xabarlariga o'z vaqtida javob "
                 "berish, belgilangan uchrashuv va suhbatlarga o'z vaqtida kelish;")
    _bullet(doc, "Tavsiya etilgan til kurslarida muntazam qatnashish;")
    _bullet(doc, "Har bir bajarilgan xizmat bo'yicha taqdim etilgan Dalolatnomani belgilangan "
                 "muddatda (3 ish kuni ichida) ko'rib chiqish va imzolash.")
    _body(doc, "7.3. Buyurtmachi yuqoridagi majburiyatlarni muntazam ravishda bajarmasa, "
                "Ijrochi yozma ogohlantirish berganidan so'ng Shartnomani bir tomonlama bekor "
                "qilish huquqiga ega, bunday holda boshlang'ich to'lov qaytarilmaydi.")

    # ---------- 8. Bekor qilish ----------
    _heading(doc, "8. SHARTNOMANI BEKOR QILISH VA TO'LOVNI QAYTARISH SHARTLARI")
    _body(doc, "8.1. Buyurtmachi Shartnomani mustaqil ravishda, muddatidan oldin (4 oylik "
                "muddat tugagunga qadar) bekor qilgan taqdirda, ilgari to'langan boshlang'ich "
                "to'lov qaytarilmaydi.")
    _body(doc, "8.2. Ish beruvchi Buyurtmachini onlayn suhbatdan o'tkazmagan yoki uning "
                "nomzodini rad etgan hollarda — ushbu holat Ijrochi faoliyatiga bog'liq "
                "bo'lmagan uchinchi shaxs (ish beruvchi)ning mustaqil qarori bo'lganligi "
                "sababli, boshlang'ich to'lov qaytarib berilmaydi, bu shartda tegishli "
                "xizmatlarning Dalolatnoma bilan tasdiqlangan holda bajarilganligi asos qilib "
                "olinadi.")
    _body(doc, "8.3. Buyurtmachi onlayn yoki bevosita suhbatdan (intervyudan) o'z aybi yoki "
                "tayyorgarligi yetarli emasligi sababli o'ta olmagan taqdirda — bu holat ham "
                "Ijrochi javobgarligini keltirib chiqarmaydi va boshlang'ich to'lov "
                "qaytarilmaydi, chunki Ijrochining majburiyati suhbatni tashkil etish va "
                "tayyorgarlik ko'rishdan iborat bo'lib, natijaning o'zi emas.")
    _body(doc, "8.4. Elchixona suhbati yoki hujjatlarni ko'rib chiqish jarayonida, hujjatlarda "
                "Ijrochi tomonidan yo'l qo'yilgan rasmiy nuqson bo'lmagan holda ham viza rad "
                "etilgan taqdirda — bu qaror to'liq elchixonaning mustaqil vakolati doirasida "
                "bo'lib, Ijrochi javobgar emas va boshlang'ich to'lov qaytarilmaydi.")
    _body(doc, "8.5. Agar Ijrochi Shartnoma muddati (4 oy) davomida 2.1-bandda sanab o'tilgan "
                "majburiyatlardan birortasini ham bajarmagan bo'lsa — bu holat tegishli "
                "Dalolatnomalarning yo'qligi yoki Ijrochi tomonidan hujjatlashtirilmagan "
                "boshqa dalillar bilan tasdiqlansa — Buyurtmachining boshlang'ich to'lovi "
                "to'liq hajmda qaytariladi.")
    _body(doc, "8.6. Agar Shartnoma bo'yicha jarayon Buyurtmachining 6.3–6.5-bandlarda "
                "nazarda tutilgan tartibda bog'lanishga chiqmaganligi, qo'ng'iroq va "
                "xabarlarga javob bermaganligi yoki taklif etilgan suhbat/uchrashuvga o'z "
                "vaqtida kelmaganligi sababli davom ettirilmagan yoki natijasiz yakunlangan "
                "bo'lsa — bu Buyurtmachining o'zi tomonidan Shartnoma bo'yicha "
                "majburiyatlarning buzilishi hisoblanadi, tegishli xizmat Aloqa jurnalida "
                "qayd etilgan urinishlar asosida bajarilgan deb tan olinadi va boshlang'ich "
                "to'lov qaytarilmaydi.")
    _body(doc, "8.7. Ushbu bandlarda ko'rsatilgan har qanday nizoli holatda, Dalolatnomalar, "
                "Aloqa jurnali, yozishmalar, qo'ng'iroqlar tarixi va boshqa 5–6-bandlarda "
                "nazarda tutilgan dalillar ustuvor ahamiyatga ega hisoblanadi.")

    # ---------- 9. Fors-major ----------
    _heading(doc, "9. FAVQULODDA VAZIYATLAR (FORS-MAJOR)")
    _body(doc, "9.1. Tabiiy ofatlar va ekologik favqulodda vaziyatlar, harbiy harakatlar yoki "
                "qurolli nizolar, terrorchilik va sabotaj, davlat organlari tomonidan "
                "belgilangan cheklov va taqiqlar, epidemiya yoki pandemiya holatlari "
                "fors-major deb tan olinadi.")
    _body(doc, "9.2. Fors-major yuzaga kelganda, ta'sirlangan tomon ikkinchi tomonga darhol "
                "(2 ish kuni ichida) yozma xabar yuborishi shart.")
    _body(doc, "9.3. Fors-major davomida majburiyatlarni bajarish kechikishi uchun Tomonlar "
                "javobgarlikdan ozod etiladi; Shartnoma muddati fors-major davomiyligi "
                "ko'lamida uzaytiriladi.")

    # ---------- 10. Nizolar ----------
    _heading(doc, "10. NIZOLARNI HAL ETISH TARTIBI")
    _body(doc, "10.1. Ushbu Shartnoma yuzasidan kelib chiqadigan barcha kelishmovchiliklar, "
                "jumladan Buyurtmachining Ijrochi tomonidan xizmat ko'rsatilmaganligi yoki "
                "sifatsiz ko'rsatilganligi to'g'risidagi har qanday da'vosi, avvalo ANIQ VA "
                "YOZMA PRETENZION TARTIBDA hal etilishi shart: Buyurtmachi o'z da'volarini "
                "yozma ravishda Ijrochiga taqdim etadi, Ijrochi esa 10 (o'n) ish kuni ichida "
                "asoslangan javob berishi shart.")
    _body(doc, "10.2. Buyurtmachi ushbu Shartnomaning 10.1-bandida nazarda tutilgan "
                "pretenzion tartibga rioya qilmasdan turib davlat organlariga (jumladan Ichki "
                "ishlar boshqarmasiga) murojaat qilgan taqdirda, Ijrochi ushbu Shartnoma, "
                "unga ilova qilingan Dalolatnomalar va boshqa barcha dalillarni tegishli "
                "organga taqdim etish orqali o'z pozitsiyasini to'liq himoya qilish huquqiga "
                "ega.")
    _body(doc, "10.3. Agar Buyurtmachining davlat organiga qilgan murojaati keyinchalik "
                "asossiz deb topilsa va bu Ijrochining nufuziga yoki faoliyatiga zarar "
                "yetkazgan bo'lsa, Ijrochi O'zbekiston Respublikasi qonunchiligida nazarda "
                "tutilgan tartibda o'z huquqlarini sud orqali himoya qilish huquqini o'zida "
                "saqlab qoladi.")
    _body(doc, "10.4. Pretenzion tartib orqali kelishuvga erishilmagan taqdirda, nizo "
                "O'zbekiston Respublikasining amaldagi qonunchiligiga muvofiq sudlarda ko'rib "
                "chiqiladi.")

    # ---------- 11. Maxfiylik ----------
    _heading(doc, "11. MAXFIYLIK")
    _body(doc, "11.1. Tomonlar Shartnoma doirasida almashilgan barcha ma'lumotlarni maxfiy "
                "saqlash majburiyatini oladilar va ularni uchinchi shaxslarga oshkor etmaydilar, "
                "ushbu Shartnomaning 6-bandida nazarda tutilgan hollar va qonun talab qilgan "
                "hollar bundan mustasno.")
    _body(doc, "11.2. Mazkur majburiyat Shartnoma muddati tugaganidan keyin ham 2 (ikki) yil "
                "davomida o'z kuchini saqlab qoladi.")

    # ---------- 12. Yakuniy qoidalar ----------
    _heading(doc, "12. YAKUNIY QOIDALAR")
    _bullet(doc, "Ushbu Shartnoma ikki nusxada — har bir Tomon uchun bittadan — tuzilgan "
                 "bo'lib, nusxalarning har biri teng yuridik kuchga ega.")
    _bullet(doc, "Shartnomaga o'zgartirish yoki qo'shimchalar kiritish faqat Tomonlarning "
                 "yozma roziligi bilan rasmiylashtiriladi.")
    _bullet(doc, "Shartnoma imzolangan paytdan e'tiboran Tomonlar uning barcha shartlari, "
                 "jumladan 3-banddagi natijaning kafolatlanmasligi to'g'risidagi shart bilan "
                 "to'liq tanishgan va unga so'zsiz rozi bo'lgan deb hisoblanadi.")
    _bullet(doc, "Tomonlar o'z huquq va majburiyatlarini uchinchi shaxslarga o'tkazish "
                 "huquqiga ega emaslar.")
    _bullet(doc, "Ushbu Shartnomaning 1-ILOVASI (Dalolatnoma namunasi) va 2-ILOVASI (Aloqa "
                 "jurnali namunasi) uning ajralmas qismi hisoblanadi.")

    # ---------- 13. Rekvizitlar va imzolar ----------
    _heading(doc, "13. TOMONLARNING REKVIZITLARI VA IMZOLARI")

    table = doc.add_table(rows=1, cols=2)
    table.columns[0].width = Cm(8)
    table.columns[1].width = Cm(8)
    hdr = table.rows[0].cells
    hdr[0].text = "IJROCHI:"
    hdr[1].text = "BUYURTMACHI:"
    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True

    rows_data = [
        (COMPANY_NAME, f"F.I.Sh.: {full_name}"),
        (f"STIR: {COMPANY_STIR}", f"Tug'ilgan kuni: {birth_date}"),
        (f"MFO: {COMPANY_MFO}", f"Pasport: {passport}"),
        (f"H/r: {COMPANY_ACCOUNT}", f"Telefon: {phone}"),
        (f"Telefon: {COMPANY_PHONE}", f"Manzil: {address}"),
        (f"Manzil: {COMPANY_ADDRESS}", ""),
        (f"Direktor: {COMPANY_DIRECTOR}", ""),
        ("Imzo: ______________________", "Imzo: ______________________"),
    ]
    for left, right in rows_data:
        row = table.add_row().cells
        row[0].text = left
        row[1].text = right

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    _body(doc, f"Sana: {today}", size=10)

    # ---------- 1-ILOVA: Dalolatnoma ----------
    doc.add_page_break()
    _heading(doc, "1-ILOVA", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    _body(doc, f"Shartnoma № {contract_number} ga ilova", size=10)
    _heading(doc, "XIZMATLAR BAJARILISHI TO'G'RISIDA DALOLATNOMA", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    _body(doc, f"Biz, quyida imzo qo'yuvchilar, {COMPANY_NAME} (Ijrochi) va {full_name} "
                f"(Buyurtmachi), № {contract_number} shartnoma bo'yicha quyidagi xizmatlarning "
                f"bajarilish holatini tasdiqlaymiz:")

    services = [
        "Mijoz hujjatlari haqiqiyligini tekshirish",
        "Mijozning sudlanmaganligi haqida ma'lumot olish",
        "MIB (aliment, kredit va h.k.) bo'yicha qarzdorligi bor-yo'qligini tekshirish",
        "O'zbekistondan chiqishga taqiq bor-yo'qligini tekshirish",
        "Hujjatlar tarjimasi (talab qilingan tilga)",
        "Ish beruvchilarga jo'natish uchun onlayn rezyume to'ldirib shakllantirish",
        "Tayyor rezyume va hujjatlarni ish beruvchiga elektron jo'natib berish",
        "Mijoz ma'qullansa, ish beruvchi bilan belgilangan suhbat kuni/vaqtidan mijozni xabardor qilish",
        "Suhbatni texnik jihatdan tashkillashtirish va o'tkazish",
        "Suhbat natijasidan mijozni xabardor qilish",
    ]
    dal_table = doc.add_table(rows=1, cols=4)
    dal_hdr = dal_table.rows[0].cells
    dal_hdr[0].text = "№"
    dal_hdr[1].text = "Xizmat nomi"
    dal_hdr[2].text = "Bajarildi"
    dal_hdr[3].text = "Sana"
    for cell in dal_hdr:
        cell.paragraphs[0].runs[0].bold = True
    for i, service in enumerate(services, 1):
        row = dal_table.add_row().cells
        row[0].text = str(i)
        row[1].text = service
        row[2].text = "☐ Ha  ☐ Yo'q"
        row[3].text = ""

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    _body(doc, "Ushbu Dalolatnoma yuqorida qayd etilgan xizmatlarning bajarilganligi (yoki "
                "bajarilmaganligi) to'g'risidagi rasmiy dalil hisoblanadi va tomonlar o'rtasida "
                "kelib chiqishi mumkin bo'lgan nizolarda asosiy hujjat sifatida ishlatiladi.")

    sign_table = doc.add_table(rows=2, cols=2)
    sign_table.rows[0].cells[0].text = "Ijrochi tomonidan: ______________________"
    sign_table.rows[0].cells[1].text = "Buyurtmachi tomonidan: ______________________"
    sign_table.rows[1].cells[0].text = "Sana: _______________"
    sign_table.rows[1].cells[1].text = "Sana: _______________"

    # ---------- 2-ILOVA: Aloqa jurnali ----------
    doc.add_page_break()
    _heading(doc, "2-ILOVA", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    _body(doc, f"Shartnoma № {contract_number} ga ilova", size=10)
    _heading(doc, "BUYURTMACHI BILAN ALOQA (QO'NG'IROQLAR VA XABARLAR) JURNALI", size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    _body(doc, "Ushbu jurnal Shartnomaning 6.4–6.6 va 8.6-bandlariga muvofiq, Ijrochi "
                "tomonidan Buyurtmachi bilan bog'lanish yuzasidan amalga oshirilgan har bir "
                "urinishni qayd etish uchun yuritiladi. Jurnaldagi yozuvlar — Buyurtmachining "
                "xabardor qilinganligi yoki bog'lanishga chiqmaganligini tasdiqlovchi rasmiy "
                "dalil hisoblanadi.")

    log_table = doc.add_table(rows=1, cols=5)
    log_hdr = log_table.rows[0].cells
    for i, h in enumerate(["№", "Sana", "Vaqt", "Aloqa usuli", "Natija"]):
        log_hdr[i].text = h
        log_hdr[i].paragraphs[0].runs[0].bold = True
    for i in range(1, 11):
        row = log_table.add_row().cells
        row[0].text = str(i)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    _body(doc, "Izoh: har bir urinish sana, aniq vaqt, qaysi usulda (qo'ng'iroq, SMS, "
                "messenjer) amalga oshirilgani va uning natijasi (javob berdi / javob "
                "bermadi / band edi / raqam o'chiq) ko'rsatilgan holda to'ldiriladi.", size=10)

    os.makedirs(output_dir, exist_ok=True)
    safe_name = full_name.replace(" ", "_").replace("/", "_")
    filepath = os.path.join(output_dir, f"shartnoma_{contract_number}_{safe_name}.docx")
    doc.save(filepath)
    return filepath
